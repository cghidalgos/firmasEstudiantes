from flask import Flask, render_template, request, redirect, url_for, flash, session, send_file
import base64
import json
import re
import subprocess
from io import BytesIO
from datetime import datetime
from uuid import uuid4
from functools import wraps

from werkzeug.utils import secure_filename

import os
from PIL import Image
from docx import Document
from docx.shared import Inches

app = Flask(__name__)

UPLOAD_FOLDER = 'firmas_temp'
ACTAS_FOLDER = 'actas'
CURSOS_FILE_LEGACY = 'cursos.txt'
CURSOS_DB = 'cursos.json'
EXAMPLE_ACTA_DOC = 'F6. Acta de inicio de curso.doc'
MANUAL_PDF = 'Manual_Profesor_Firmas.pdf'

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(ACTAS_FOLDER, exist_ok=True)
os.makedirs(os.path.join(ACTAS_FOLDER, '_ejemplos'), exist_ok=True)

# Config básica
app.secret_key = os.environ.get('SECRET_KEY', 'dev-insecure-secret-key')

# Credenciales simples (según requerimiento)
ADMIN_USER = os.environ.get('ADMIN_USER', 'admin')
ADMIN_PASS = os.environ.get('ADMIN_PASS', 'admin2026')
DELETE_COURSE_PASS = os.environ.get('DELETE_COURSE_PASS', 'lidis2026*')


@app.context_processor
def inject_globals():
    return {
        'current_year': datetime.now().year,
    }


def _load_courses_db():
    """Retorna lista de cursos: [{nombre, profesor, acta_path}]."""
    if not os.path.exists(CURSOS_DB):
        return []
    try:
        with open(CURSOS_DB, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if isinstance(data, list):
                return data
    except Exception:
        return []
    return []


def _save_courses_db(courses):
    tmp_path = f"{CURSOS_DB}.tmp"
    with open(tmp_path, 'w', encoding='utf-8') as f:
        json.dump(courses, f, ensure_ascii=False, indent=2)
    os.replace(tmp_path, CURSOS_DB)


def _normalize_course_name(name: str) -> str:
    return re.sub(r"\s+", " ", (name or "").strip())


def require_admin(view_func):
    @wraps(view_func)
    def wrapper(*args, **kwargs):
        if not session.get('is_admin'):
            flash('Inicia sesión para acceder a administración.', 'warning')
            return redirect(url_for('login', next=request.path))
        return view_func(*args, **kwargs)
    return wrapper


def get_courses():
    """Lista cursos; si no hay DB nueva, usa el archivo legacy cursos.txt."""
    courses = _load_courses_db()
    if courses:
        # Normalizar estructura: asegurar id
        changed = False
        for c in courses:
            if 'id' not in c or not c.get('id'):
                c['id'] = uuid4().hex
                changed = True
        if changed:
            _save_courses_db(courses)
        return courses

    # Legacy
    if not os.path.exists(CURSOS_FILE_LEGACY):
        return []
    with open(CURSOS_FILE_LEGACY, 'r', encoding='utf-8') as f:
        return [{
            'id': uuid4().hex,
            'nombre': line.strip(),
            'profesor': '',
            'acta_path': '',
        } for line in f if line.strip()]


def get_cursos():
    return [c.get('nombre', '') for c in get_courses() if c.get('nombre')]


def get_course_by_name(course_name: str):
    course_name = _normalize_course_name(course_name)
    for c in get_courses():
        if _normalize_course_name(c.get('nombre', '')) == course_name:
            return c
    return None


def get_course_by_id(course_id: str):
    for c in get_courses():
        if str(c.get('id')) == str(course_id):
            return c
    return None


def _course_acta_default_path(course_name: str) -> str:
    # Requisito: f6_nombre del curso.docx (sanitizado)
    safe = secure_filename(course_name).strip('_')
    if not safe:
        safe = 'curso'
    return os.path.join(ACTAS_FOLDER, f"f6_{safe}.docx")


def _soffice_available() -> bool:
    try:
        res = subprocess.run(['soffice', '--version'], capture_output=True, text=True)
        return res.returncode == 0
    except Exception:
        return False


def _convert_doc_to_docx(doc_path: str, target_docx_path: str) -> None:
    """Convierte un .doc a .docx usando LibreOffice headless.

    Crea el .docx dentro de ACTAS_FOLDER y lo renombra a target_docx_path.
    """
    if not _soffice_available():
        raise RuntimeError('LibreOffice (soffice) no está disponible para convertir .doc a .docx')

    out_dir = os.path.dirname(os.path.abspath(target_docx_path))
    os.makedirs(out_dir, exist_ok=True)

    # LibreOffice genera <base>.docx en outdir.
    cmd = ['soffice', '--headless', '--nologo', '--nolockcheck', '--nodefault', '--norestore', '--convert-to', 'docx', '--outdir', out_dir, os.path.abspath(doc_path)]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"Error convirtiendo .doc a .docx: {result.stderr or result.stdout}")

    base = os.path.splitext(os.path.basename(doc_path))[0]
    generated = os.path.join(out_dir, f"{base}.docx")
    if not os.path.exists(generated):
        raise RuntimeError('No se encontró el .docx generado por LibreOffice')

    os.replace(generated, target_docx_path)


@app.route('/admin/acta-ejemplo')
@require_admin
def download_example_acta():
    fmt = (request.args.get('format') or 'docx').lower()
    example_doc_path = os.path.join(app.root_path, EXAMPLE_ACTA_DOC)
    if not os.path.exists(example_doc_path):
        flash('No se encontró el archivo de ejemplo en el servidor.', 'error')
        return redirect(url_for('crear_curso'))

    if fmt == 'doc' or fmt == 'docx':
        if fmt == 'docx':
            flash('El ejemplo solo está disponible como .doc.', 'info')
        return send_file(
            os.path.abspath(example_doc_path),
            as_attachment=True,
            download_name=os.path.basename(example_doc_path),
        )

    flash('Formato no válido. Usa format=doc.', 'error')
    return redirect(url_for('crear_curso'))


@app.route('/admin/manual-usuario')
@require_admin
def admin_manual_usuario():
    manual_pdf_path = os.path.join(app.root_path, MANUAL_PDF)
    if not os.path.exists(manual_pdf_path):
        flash('No se encontró el manual de usuario en el servidor.', 'error')
        return redirect(url_for('crear_curso'))

    return send_file(
        os.path.abspath(manual_pdf_path),
        mimetype='application/pdf',
        as_attachment=False,
        download_name=os.path.basename(manual_pdf_path),
    )


def _course_acta_filename(course_name: str) -> str:
    return os.path.basename(_course_acta_default_path(course_name))


def find_table_with_headers(doc):
    """Buscar una tabla que contenga encabezados relacionados con Nombre, Codigo y Firma.
    Devuelve (table, name_idx, code_idx, firma_idx) o (None, None, None, None).
    """
    for table in doc.tables:
        # leer textos de la primera fila
        if len(table.rows) == 0:
            continue
        headers = [cell.text.strip().upper().replace(' ', '') for cell in table.rows[0].cells]
        # buscar coincidencias parciales
        has_name = any('NOMBRE' in h or 'NOMBRECOMPLETO' in h for h in headers)
        has_code = any('CODIG' in h for h in headers)
        has_firma = any('FIRMA' in h for h in headers)
        if has_name and has_code and has_firma:
            # obtener índices
            name_idx = next((i for i,h in enumerate(headers) if 'NOMBRE' in h or 'NOMBRECOMPLETO' in h), 0)
            code_idx = next((i for i,h in enumerate(headers) if 'CODIG' in h), 1)
            firma_idx = next((i for i,h in enumerate(headers) if 'FIRMA' in h), 2)
            return table, name_idx, code_idx, firma_idx
    return None, None, None, None


def resolve_word_path_for_course(curso: str) -> str:
    """Resuelve la ruta del acta (.docx) para un curso.

    Prioridad:
    1) `acta_path` guardado en cursos.json
    2) fallback legacy en actas/ o raíz
    3) ruta por defecto en actas/
    """
    curso = _normalize_course_name(curso)
    course = get_course_by_name(curso)
    if course and course.get('acta_path') and os.path.exists(course.get('acta_path')):
        return course.get('acta_path')

    fallback1 = _course_acta_default_path(curso).replace('f6_', 'F6_Acta_')
    fallback2 = f"F6_Acta_{curso.replace(' ', '_')}.docx"
    if os.path.exists(fallback1):
        return fallback1
    if os.path.exists(fallback2):
        return fallback2
    return _course_acta_default_path(curso)


def has_codigo_already_signed(doc: Document, codigo: str) -> bool:
    codigo = _normalize_course_name(codigo)
    table, _, code_idx, _ = find_table_with_headers(doc)
    if table is None:
        return False
    for row in table.rows[1:]:
        if len(row.cells) > code_idx:
            if _normalize_course_name(row.cells[code_idx].text) == codigo:
                return True
    return False


@app.route('/api/firmas/existe', methods=['GET'])
def api_firma_existe():
    curso = _normalize_course_name(request.args.get('curso', ''))
    codigo = _normalize_course_name(request.args.get('codigo', ''))
    if not curso or not codigo:
        return {'ok': False, 'error': 'curso y codigo son requeridos'}, 400

    word_path = resolve_word_path_for_course(curso)
    if not os.path.exists(word_path):
        # Si no hay acta aún, nadie ha firmado
        return {'ok': True, 'exists': False}

    try:
        doc = Document(word_path)
        exists = has_codigo_already_signed(doc, codigo)
        return {'ok': True, 'exists': bool(exists)}
    except Exception:
        # Si falla leer el docx, no bloquear al usuario en frontend
        return {'ok': True, 'exists': False}

@app.route('/')
def index():
    cursos = get_courses()
    return render_template('form.html', cursos=cursos)


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = (request.form.get('username') or '').strip()
        password = (request.form.get('password') or '').strip()
        if username == ADMIN_USER and password == ADMIN_PASS:
            session['is_admin'] = True
            flash('Sesión iniciada.', 'success')
            next_url = request.args.get('next')
            return redirect(next_url or url_for('crear_curso'))
        flash('Credenciales incorrectas.', 'error')
        return redirect(url_for('login'))
    return render_template('login.html')


@app.route('/logout')
def logout():
    session.pop('is_admin', None)
    flash('Sesión cerrada.', 'success')
    return redirect(url_for('index'))


@app.route('/gracias')
def gracias():
    last = session.pop('last_submission', None) or {}
    return render_template('gracias.html', **last)

@app.route('/crear_curso', methods=['GET', 'POST'])
@require_admin
def crear_curso():
    if request.method == 'POST':
        nombre_curso = _normalize_course_name(request.form.get('nombre_curso', ''))
        profesor = _normalize_course_name(request.form.get('profesor', ''))
        acta_file = request.files.get('acta')

        if not profesor:
            flash('Por favor ingresa el nombre del profesor.', 'error')
            return redirect(url_for('crear_curso'))
        if not nombre_curso:
            flash('Por favor ingresa el nombre del curso.', 'error')
            return redirect(url_for('crear_curso'))
        if not acta_file or not acta_file.filename:
            flash('Por favor sube el acta en formato .docx.', 'error')
            return redirect(url_for('crear_curso'))

        filename = acta_file.filename
        lower = filename.lower()
        if not (lower.endswith('.docx') or lower.endswith('.doc')):
            flash('El acta debe ser un archivo .docx o .doc.', 'error')
            return redirect(url_for('crear_curso'))

        acta_path = _course_acta_default_path(nombre_curso)  # siempre .docx
        tmp_dir = os.path.join(ACTAS_FOLDER, '_tmp')
        os.makedirs(tmp_dir, exist_ok=True)

        try:
            if lower.endswith('.docx'):
                acta_file.save(acta_path)
            else:
                # Guardar .doc temporal y convertirlo a .docx
                tmp_doc = os.path.join(tmp_dir, f"upload_{uuid4().hex}.doc")
                acta_file.save(tmp_doc)
                _convert_doc_to_docx(tmp_doc, acta_path)
                try:
                    os.remove(tmp_doc)
                except Exception:
                    pass

            # Validar que el resultado final sea un docx legible
            Document(acta_path)
        except Exception:
            try:
                if os.path.exists(acta_path):
                    os.remove(acta_path)
            except Exception:
                pass
            flash('No se pudo procesar el archivo. Asegúrate de que sea un .docx válido o un .doc convertible.', 'error')
            return redirect(url_for('crear_curso'))

        courses = _load_courses_db()
        # Si no existe DB y hay legacy, migrar cursos legacy sin acta
        if not courses and os.path.exists(CURSOS_FILE_LEGACY):
            for legacy_name in get_cursos():
                courses.append({'id': uuid4().hex, 'nombre': legacy_name, 'profesor': '', 'acta_path': ''})

        # Evitar duplicados por nombre
        existing_idx = next((i for i, c in enumerate(courses)
                             if _normalize_course_name(c.get('nombre', '')) == nombre_curso), None)
        new_course = {'id': uuid4().hex, 'nombre': nombre_curso, 'profesor': profesor, 'acta_path': acta_path}
        if existing_idx is None:
            courses.append(new_course)
        else:
            # Mantener id existente
            keep_id = courses[existing_idx].get('id') or uuid4().hex
            courses[existing_idx] = {**courses[existing_idx], **new_course, 'id': keep_id}

        _save_courses_db(courses)
        flash('Curso creado. Los estudiantes ya pueden firmar el acta.', 'success')
        return redirect(url_for('crear_curso'))

    cursos = get_courses()
    return render_template('crear_curso.html', cursos=cursos)


@app.route('/cursos/<course_id>/editar', methods=['GET', 'POST'])
@require_admin
def editar_curso(course_id):
    course = get_course_by_id(course_id)
    if not course:
        flash('Curso no encontrado.', 'error')
        return redirect(url_for('crear_curso'))

    if request.method == 'POST':
        new_name = _normalize_course_name(request.form.get('nombre_curso', ''))
        new_prof = _normalize_course_name(request.form.get('profesor', ''))
        acta_file = request.files.get('acta')

        if not new_prof:
            flash('Por favor ingresa el nombre del profesor.', 'error')
            return redirect(url_for('editar_curso', course_id=course_id))
        if not new_name:
            flash('Por favor ingresa el nombre del curso.', 'error')
            return redirect(url_for('editar_curso', course_id=course_id))

        courses = _load_courses_db() or get_courses()
        # Validar duplicado por nombre (excepto el mismo curso)
        for c in courses:
            if str(c.get('id')) != str(course_id) and _normalize_course_name(c.get('nombre', '')) == new_name:
                flash('Ya existe un curso con ese nombre.', 'error')
                return redirect(url_for('editar_curso', course_id=course_id))

        target_acta_path = _course_acta_default_path(new_name)

        # Si suben un acta nueva, guardarla/convertirla en el nombre objetivo (siempre .docx)
        if acta_file and acta_file.filename:
            lower = acta_file.filename.lower()
            if not (lower.endswith('.docx') or lower.endswith('.doc')):
                flash('El acta debe ser un archivo .docx o .doc.', 'error')
                return redirect(url_for('editar_curso', course_id=course_id))
            try:
                if lower.endswith('.docx'):
                    acta_file.save(target_acta_path)
                else:
                    tmp_dir = os.path.join(ACTAS_FOLDER, '_tmp')
                    os.makedirs(tmp_dir, exist_ok=True)
                    tmp_doc = os.path.join(tmp_dir, f"upload_{uuid4().hex}.doc")
                    acta_file.save(tmp_doc)
                    _convert_doc_to_docx(tmp_doc, target_acta_path)
                    try:
                        os.remove(tmp_doc)
                    except Exception:
                        pass

                Document(target_acta_path)
            except Exception:
                try:
                    if os.path.exists(target_acta_path):
                        os.remove(target_acta_path)
                except Exception:
                    pass
                flash('No se pudo procesar el archivo. Asegúrate de que sea un .docx válido o un .doc convertible.', 'error')
                return redirect(url_for('editar_curso', course_id=course_id))
        else:
            # Si no suben acta nueva y cambian el nombre, renombrar el archivo existente si existe
            old_path = course.get('acta_path') or ''
            if old_path and os.path.exists(old_path) and os.path.abspath(old_path) != os.path.abspath(target_acta_path):
                try:
                    os.replace(old_path, target_acta_path)
                except Exception:
                    # Si falla renombrar, dejar el path actual
                    target_acta_path = old_path

        # Persistir cambios
        updated = False
        for i, c in enumerate(courses):
            if str(c.get('id')) == str(course_id):
                courses[i] = {
                    **c,
                    'nombre': new_name,
                    'profesor': new_prof,
                    'acta_path': target_acta_path,
                }
                updated = True
                break
        if updated:
            _save_courses_db(courses)
            flash('Curso actualizado.', 'success')
        return redirect(url_for('editar_curso', course_id=course_id))

    view = {
        'id': course.get('id'),
        'nombre': course.get('nombre', ''),
        'profesor': course.get('profesor', ''),
        'acta_filename': os.path.basename(course.get('acta_path') or _course_acta_default_path(course.get('nombre', ''))),
    }
    return render_template('editar_curso.html', course=view)


@app.route('/cursos/<course_id>/acta')
@require_admin
def download_acta(course_id):
    course = get_course_by_id(course_id)
    if not course:
        flash('Curso no encontrado.', 'error')
        return redirect(url_for('crear_curso'))
    path = course.get('acta_path') or _course_acta_default_path(course.get('nombre', ''))
    if not path or not os.path.exists(path):
        flash('El acta no existe para este curso.', 'error')
        return redirect(url_for('editar_curso', course_id=course_id))

    abs_path = os.path.abspath(path)
    abs_actas = os.path.abspath(ACTAS_FOLDER)
    if not abs_path.startswith(abs_actas + os.sep) and abs_path != abs_actas:
        flash('Ruta de acta inválida.', 'error')
        return redirect(url_for('crear_curso'))

    return send_file(abs_path, as_attachment=True, download_name=os.path.basename(abs_path))


@app.route('/cursos/<course_id>/eliminar', methods=['POST'])
@require_admin
def eliminar_curso(course_id):
    fallback_redirect = request.referrer or url_for('crear_curso')
    provided = (request.form.get('delete_password') or '').strip()
    if provided != DELETE_COURSE_PASS:
        flash('Contraseña de eliminación incorrecta.', 'error')
        return redirect(fallback_redirect)

    courses = _load_courses_db() or []
    idx = next((i for i, c in enumerate(courses) if str(c.get('id')) == str(course_id)), None)
    if idx is None:
        flash('Curso no encontrado.', 'error')
        return redirect(fallback_redirect)

    course = courses[idx]
    acta_path = course.get('acta_path') or ''

    # Eliminar registro del curso
    courses.pop(idx)
    _save_courses_db(courses)

    # Borrar el acta solo si está dentro de actas/ (evita borrar rutas inesperadas)
    try:
        if acta_path and os.path.exists(acta_path):
            abs_path = os.path.abspath(acta_path)
            abs_actas = os.path.abspath(ACTAS_FOLDER)
            if abs_path.startswith(abs_actas + os.sep):
                os.remove(abs_path)
    except Exception:
        # No bloquear la eliminación del curso si falla el borrado del archivo
        pass

    flash('Curso eliminado.', 'success')
    return redirect(url_for('crear_curso'))


@app.route('/submit', methods=['POST'])
def submit():
    curso = _normalize_course_name(request.form.get('curso', ''))
    nombre = _normalize_course_name(request.form.get('nombre', ''))
    codigo = _normalize_course_name(request.form.get('codigo', ''))
    firma_data = (request.form.get('firma', '') or '').strip()

    if not curso:
        flash('Selecciona un curso válido.', 'error')
        return redirect(url_for('index'))
    if not nombre or len(nombre) < 3:
        flash('Ingresa tu nombre completo.', 'error')
        return redirect(url_for('index'))
    if not codigo or len(codigo) < 3:
        flash('Ingresa un código válido.', 'error')
        return redirect(url_for('index'))

    word_path = resolve_word_path_for_course(curso)

    # Cargar documento o crearlo si no existe
    if os.path.exists(word_path):
        doc = Document(word_path)
    else:
        doc = Document()
        doc.add_heading(f"Acta F6 - {curso}", 0)

    table, name_idx, code_idx, firma_idx = find_table_with_headers(doc)

    # Si no existe la tabla con los encabezados, crearla al final
    if table is None:
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'NOMBRE COMPLETO'
        hdr_cells[1].text = 'CODIGO'
        hdr_cells[2].text = 'FIRMA'
        name_idx, code_idx, firma_idx = 0, 1, 2

    if has_codigo_already_signed(doc, codigo):
        flash(f"El código {codigo} ya está registrado en el curso {curso}.", 'error')
        return redirect(url_for('index'))

    # Procesar imagen de la firma
    if ',' not in firma_data or not firma_data.startswith('data:image'):
        flash('Firma inválida. Por favor vuelve a firmar.', 'error')
        return redirect(url_for('index'))

    try:
        _, encoded = firma_data.split(',', 1)
        img_bytes = base64.b64decode(encoded)
        img = Image.open(BytesIO(img_bytes))
    except Exception:
        flash('No se pudo leer la firma. Por favor vuelve a intentarlo.', 'error')
        return redirect(url_for('index'))

    # Asegurar que la firma quede en negro en el documento (sin importar el color dibujado)
    try:
        if img.mode in ('RGBA', 'LA') or ('transparency' in img.info):
            img = img.convert('RGBA')
            r, g, b, a = img.split()
            black_rgb = Image.new('RGB', img.size, (0, 0, 0))
            br, bg, bb = black_rgb.split()
            img = Image.merge('RGBA', (br, bg, bb, a))
        else:
            # Si viene sin canal alpha (poco común desde canvas), hacer blanco->transparente y todo lo demás negro
            rgb = img.convert('RGB')
            gray = rgb.convert('L')
            # pixels claros (casi blanco) => fondo
            alpha = gray.point(lambda p: 0 if p > 245 else 255)
            black_rgb = Image.new('RGB', rgb.size, (0, 0, 0))
            br, bg, bb = black_rgb.split()
            img = Image.merge('RGBA', (br, bg, bb, alpha))
    except Exception:
        # Si algo falla en la normalización de color, continuar con la imagen original
        pass

    safe_course = secure_filename(curso).strip('_') or 'curso'
    safe_code = secure_filename(codigo).strip('_') or 'codigo'
    unique = uuid4().hex[:10]
    img_path = os.path.join(UPLOAD_FOLDER, f"{safe_course}_{safe_code}_{unique}.png")
    img.save(img_path)

    # Agregar nueva fila usando los índices de columnas encontrados
    new_cells = table.add_row().cells
    # Asegurar que la fila tenga al menos 3 celdas
    while len(new_cells) < 3:
        new_cells.append(new_cells[-1])
    new_cells[name_idx].text = nombre
    new_cells[code_idx].text = codigo
    # Insertar imagen en la celda de firma
    run = new_cells[firma_idx].paragraphs[0].add_run()
    run.add_picture(img_path, width=Inches(1.5))
    doc.save(word_path)

    session['last_submission'] = {
        'curso': curso,
        'nombre': nombre,
        'codigo': codigo,
    }
    flash('Registro guardado.', 'success')
    return redirect(url_for('gracias'))

if __name__ == '__main__':
    debug = os.environ.get('FLASK_DEBUG', '0') in ('1', 'true', 'True')
    port = int(os.environ.get('PORT', '5006'))
    app.run(debug=debug, host='0.0.0.0', port=port)